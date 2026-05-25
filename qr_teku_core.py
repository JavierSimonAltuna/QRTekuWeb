"""
QR Teku · Lógica de negocio (v3 con diseño Word PRO003)
========================================================
Funciones puras Python — sin Tkinter ni UI. Las invoca api.py desde el JS.

Configura las constantes ODBC y rutas al inicio según tu entorno.
"""

from __future__ import annotations

import io
import os
import re
import json
import platform
import subprocess
import shutil
from pathlib import Path
from datetime import datetime

import pandas as pd
import qrcode

# ─── Configuración (ajusta a tu entorno real) ───────────────────────────
ODBC_DSN = "INFOLOG"
ODBC_UID = "infolog"
ODBC_PWD = "infolog"
TABLE_NAME = "FGE50STO.GEZCAM"

TABLE_CHF_PATH = r"C:\Users\QrCarga\TABLA CHF.xlsx"

# Carpeta de guardado de los Word generados.
# Path.cwd() falla cuando la app se ejecuta dentro de OneDrive (carpeta sincronizada
# bloquea los archivos y devuelve "Bad file descriptor"). Usamos %LOCALAPPDATA% como
# destino seguro fuera de OneDrive.
def _resolve_save_dir() -> Path:
    import sys, os
    # 1) Override explícito por variable de entorno
    env = os.environ.get("QRTEKU_SAVE_DIR")
    if env:
        return Path(env)
    # 2) Si el .exe está frozen Y NO en OneDrive, junto al .exe
    if getattr(sys, "frozen", False):
        exe_dir = Path(sys.executable).resolve().parent
        if "OneDrive" not in str(exe_dir):
            return exe_dir / "QR_WORDS"
    # 3) %LOCALAPPDATA%\QRTeku\QR_WORDS  (Windows, fuera de OneDrive)
    local = os.environ.get("LOCALAPPDATA")
    if local:
        return Path(local) / "QRTeku" / "QR_WORDS"
    # 4) Documentos del usuario
    return Path.home() / "Documents" / "QRTeku" / "QR_WORDS"

SAVE_DIR = _resolve_save_dir()

# Caché ODBC en memoria + tabla CHF Excel (fallback)
_odbc_cache: dict[str, tuple[str, str]] = {}
_df_chf_cache: pd.DataFrame | None = None


# ─── Utilidades ─────────────────────────────────────────────────────────
def _norm_tractor(s: str) -> str:
    return (s or "").upper().replace(" ", "").replace("-", "").strip()


def _safe_str(v) -> str:
    if v is None:
        return ""
    try:
        if pd.isna(v):
            return ""
    except Exception:
        pass
    s = str(v).strip()
    return "" if s.lower() == "nan" else s


def _norm_hora(v) -> str:
    """Normaliza un valor de hora a 'HH:MM'.
    Acepta: 'H:MM', 'HH:MM', 'HH:MM:SS', 'YYYY-MM-DD HH:MM:SS',
    datetime.time, datetime.datetime, pd.Timestamp,
    o un float fracción-de-día de Excel (0.0833 → 02:00).
    Devuelve '' si no se reconoce."""
    if v is None:
        return ""
    # Pandas / numpy NaN
    try:
        if pd.isna(v):
            return ""
    except Exception:
        pass
    # Tipos datetime nativos + pandas Timestamp
    try:
        import datetime as _dt
        if isinstance(v, _dt.time):
            return f"{v.hour:02d}:{v.minute:02d}"
        if isinstance(v, _dt.datetime):
            return f"{v.hour:02d}:{v.minute:02d}"
        if hasattr(v, "hour") and hasattr(v, "minute"):  # pd.Timestamp etc.
            return f"{int(v.hour):02d}:{int(v.minute):02d}"
    except Exception:
        pass
    # Número float (Excel: 1.0 = 24h, 0.0833 ≈ 02:00)
    try:
        if isinstance(v, (int, float)) and not isinstance(v, bool):
            f = float(v)
            if 0 <= f < 1.0001:
                total_min = int(round(f * 24 * 60))
                return f"{(total_min // 60) % 24:02d}:{total_min % 60:02d}"
    except Exception:
        pass
    # String — soporta "HH:MM", "HH:MM:SS", "YYYY-MM-DD HH:MM[:SS]"
    s = str(v).strip()
    if not s or s.lower() == "nan":
        return ""
    # Si trae fecha+hora, nos quedamos con la parte tras el espacio
    if " " in s and ":" in s.split(" ")[-1]:
        s = s.split(" ")[-1]
    if ":" in s:
        parts = s.split(":")
        try:
            h = int(parts[0]); m = int(parts[1])
            return f"{h % 24:02d}:{m:02d}"
        except (ValueError, IndexError):
            return ""
    return ""


def _slug(s: str) -> str:
    s = (s or "").strip()
    if not s:
        return ""
    s = re.sub(r'[\\/:*?"<>|]+', "_", s)
    s = re.sub(r"\s+", "_", s)
    return s[:60]


def _make_unique_columns(cols):
    seen = {}
    new = []
    for c in cols:
        base = str(c)
        if base not in seen:
            seen[base] = 1
            new.append(base)
        else:
            seen[base] += 1
            new.append(f"{base} ({seen[base]})")
    return new


# ─── Tabla CHF Excel (fallback de ODBC) ─────────────────────────────────
def _load_chf_table() -> pd.DataFrame:
    """Carga la tabla CHF de Excel como fallback. Cachea el resultado."""
    global _df_chf_cache
    if _df_chf_cache is not None:
        return _df_chf_cache
    try:
        ext = os.path.splitext(TABLE_CHF_PATH)[1].lower()
        engine = "xlrd" if ext == ".xls" else "openpyxl"
        df = pd.read_excel(TABLE_CHF_PATH, dtype=str, engine=engine)
        df.columns = [str(c).strip() for c in df.columns]

        # Detectar columna AGENCIA (C en el Excel) y CIF (E en el Excel)
        c_col = e_col = None
        for col in df.columns:
            name = str(col).strip().lower()
            if ("agencia" in name or name == "c" or "ige" in name) and c_col is None:
                c_col = col
            if ("cif" in name or name == "e") and e_col is None:
                e_col = col

        key_col = df.columns[0]
        if c_col is None and len(df.columns) > 1:
            c_col = df.columns[1]
        if e_col is None and len(df.columns) > 2:
            e_col = df.columns[2]

        df["TRACTORA"] = df[key_col].astype(str).map(_norm_tractor)
        df["AGENCIA"] = df[c_col].apply(_safe_str) if c_col in df.columns else ""
        df["CIF"] = df[e_col].apply(_safe_str) if e_col in df.columns else ""
        _df_chf_cache = df[["TRACTORA", "AGENCIA", "CIF"]]
    except Exception:
        _df_chf_cache = pd.DataFrame(columns=["TRACTORA", "AGENCIA", "CIF"])
    return _df_chf_cache


# ─── Excel ──────────────────────────────────────────────────────────────
def load_excel(path: str) -> tuple[list[dict], str]:
    """
    Carga el Excel y devuelve (rows, fecha_b2).
    Cada row es un dict con las claves que espera el frontend.

    Para evitar 'Permission denied' cuando el Excel está abierto por otra app,
    copiamos el archivo a una ruta temporal antes de leerlo.
    """
    import tempfile, shutil
    # Copia segura: si el archivo está bloqueado por Excel, shutil.copy2 aún
    # puede leerlo con SHARE_READ y crear una copia que podemos abrir libremente.
    try:
        tmp_path = Path(tempfile.gettempdir()) / f"_qrteku_{os.getpid()}_{Path(path).name}"
        shutil.copy2(path, str(tmp_path))
        read_path = str(tmp_path)
    except Exception:
        # Fallback: intentar abrir el original directamente
        read_path = path
        tmp_path = None

    try:
        if read_path.lower().endswith(".csv"):
            raw = pd.read_csv(read_path, header=None, dtype=str)
        else:
            raw = pd.read_excel(read_path, header=None, dtype=str)
    finally:
        # Limpiar la copia temporal
        if tmp_path is not None:
            try: tmp_path.unlink(missing_ok=True)
            except Exception: pass

    # Fecha B2
    fecha_b2 = ""
    try:
        v = raw.iloc[1, 1] if len(raw) > 1 and len(raw.columns) > 1 else ""
        if pd.notna(v) and str(v).strip():
            if isinstance(v, pd.Timestamp):
                fecha_b2 = v.strftime("%Y%m%d")
            else:
                s = str(v).strip()
                try:
                    if re.match(r"^\d{4}-\d{2}-\d{2}", s):
                        fecha_b2 = pd.to_datetime(s).strftime("%Y%m%d")
                    else:
                        fecha_b2 = pd.to_datetime(s, dayfirst=True).strftime("%Y%m%d")
                except Exception:
                    fecha_b2 = s if (s.isdigit() and len(s) == 8) else pd.Timestamp.now().strftime("%Y%m%d")
        else:
            fecha_b2 = pd.Timestamp.now().strftime("%Y%m%d")
    except Exception:
        fecha_b2 = pd.Timestamp.now().strftime("%Y%m%d")

    # Encabezados (fila con 'DESTINO')
    hdr = 0
    for i in range(min(50, len(raw))):
        if "DESTINO" in raw.iloc[i].astype(str).values:
            hdr = i
            break

    raw.columns = _make_unique_columns(raw.iloc[hdr])
    df = raw.iloc[hdr + 1:].reset_index(drop=True).head(400)

    # Columna precinto (cualquier que contenga PRECINTO, o índice 30 = AE)
    precinto_col = None
    for c in df.columns:
        if "PRECINTO" in str(c).upper():
            precinto_col = c
            break
    if precinto_col is None and len(df.columns) > 30:
        precinto_col = df.columns[30]

    # Para evitar dejar matrículas iguales sin precinto distinto, agrupamos
    # los precintos por Nº viaje para cada fila.
    n_col = "Nº" if "Nº" in df.columns else ("N°" if "N°" in df.columns else None)

    def _norm_n(x):
        s = _safe_str(x).strip()
        return s.lstrip("0") if s.lstrip("0") else s

    rows: list[dict] = []
    for _, r in df.iterrows():
        destino = _safe_str(r.get("DESTINO", ""))
        n = _safe_str(r.get("Nº", r.get("N°", "")))
        agencia = _safe_str(r.get("AGENCIA CONTRATADA", ""))
        matriculas = _safe_str(r.get("MATRICULAS", ""))
        tipo = _safe_str(r.get("TIPO CAMION", ""))
        expedicion = _safe_str(r.get("EXPEDICION", ""))
        cad = _safe_str(r.get("C.A.D", ""))
        orden = _safe_str(r.get("ORDEN CARGA OPERADOR", ""))
        # PLAYA (col K) y MUELLE (col L)
        playa = ""
        muelle = ""
        for col in df.columns:
            cu = str(col).upper().replace("\n", "").replace(" ", "").replace("\t", "")
            if not playa and cu == "PLAYA":
                playa = _safe_str(r.get(col, ""))
            if not muelle and cu == "MUELLE":
                muelle = _safe_str(r.get(col, ""))
        if not playa and len(df.columns) > 10:
            playa = _safe_str(r.iloc[10])
        if not muelle and len(df.columns) > 11:
            muelle = _safe_str(r.iloc[11])
        # HORA ACULE (columna AA = índice 26)
        hora_acule = ""
        for col in df.columns:
            cu = str(col).upper().replace("\n", "").replace(" ", "").replace("\t", "")
            if not hora_acule and "ACULE" in cu:
                hora_acule = _norm_hora(r.get(col, ""))
                break
        if not hora_acule and len(df.columns) > 26:
            hora_acule = _norm_hora(r.iloc[26])
        aculado = bool(hora_acule.strip())

        # HORA SALIDA PREVISTA (columna AB = índice 27, etiqueta "SALIDA PREV")
        hora_salida = ""
        for col in df.columns:
            cu = str(col).upper().replace("\n", "").replace(" ", "").replace("\t", "")
            if not hora_salida and ("SALIDAPREV" in cu or "SALIDA_PREV" in cu or cu == "SALIDA"):
                hora_salida = _norm_hora(r.get(col, ""))
                break
        if not hora_salida and len(df.columns) > 27:
            hora_salida = _norm_hora(r.iloc[27])

        # Recoger todos los precintos con el mismo Nº viaje (cada uno con su centro)
        precintos_data = []
        if precinto_col is not None and n_col and n:
            target = _norm_n(n)
            try:
                matching = df[df[n_col].apply(_norm_n) == target]
                seen = set()
                for _, mrow in matching.iterrows():
                    pv = _safe_str(mrow.get(precinto_col, ""))
                    if pv and pv not in seen:
                        seen.add(pv)
                        centro_v = _safe_str(mrow.get("DESTINO", ""))
                        precintos_data.append({"centro": centro_v, "precinto": pv})
            except Exception:
                pass
        if not precintos_data and precinto_col is not None:
            only = _safe_str(r.get(precinto_col, ""))
            if only:
                precintos_data = [{"centro": destino, "precinto": only}]

        # Viaje combinado (precintos con distintos centros) → invertir el orden,
        # porque el orden de carga es inverso al orden de descarga del Excel.
        _centros_distintos = {(_safe_str(p.get("centro", "")).upper()) for p in precintos_data if _safe_str(p.get("centro", ""))}
        if len(_centros_distintos) >= 2:
            precintos_data = list(reversed(precintos_data))

        precinto = ",".join(p["precinto"] for p in precintos_data)

        rows.append({
            "destino": destino, "n": n, "agencia": agencia,
            "matriculas": matriculas, "tipo": tipo, "expedicion": expedicion,
            "cad": cad, "orden": orden,
            "playa": playa, "muelle": muelle,
            "hora_acule": hora_acule, "aculado": aculado,
            "hora_salida": hora_salida,
            "precinto": precinto,
            "precintos_data": precintos_data,
            "estado": "ready" if destino else "missing-cif",
        })

    # Quitar duplicados con el mismo Nº viaje (mantener solo la primera)
    seen_n = set()
    deduped = []
    for row in rows:
        key = (row["destino"], row["n"])
        if row["n"] and key in seen_n:
            continue
        seen_n.add(key)
        deduped.append(row)

    return deduped, fecha_b2


# ─── ODBC + CHF Excel fallback ──────────────────────────────────────────
def _sql_norm(col: str) -> str:
    return f"REPLACE(REPLACE(UPPER({col}), ' ', ''), '-', '')"


def odbc_lookup_chf(matricula: str) -> tuple[str, str]:
    """
    Busca CIF + Agencia. Primero por ODBC, luego por Excel CHF (fallback).
    Devuelve (CIF, Agencia). Cachea resultados.
    """
    tractora_norm = _norm_tractor(matricula)
    if not tractora_norm:
        return "", ""
    if tractora_norm in _odbc_cache:
        return _odbc_cache[tractora_norm]

    cif, agencia = "", ""

    # Intento 1: ODBC
    try:
        import pyodbc
        conn = pyodbc.connect(
            f"DSN={ODBC_DSN};UID={ODBC_UID};PWD={ODBC_PWD};",
            timeout=8, autocommit=True,
        )
        try:
            cur = conn.cursor()
            sql_norm = _sql_norm("CODCAM")
            queries = [
                (f"SELECT CODTRA, CODCHF FROM {TABLE_NAME} WHERE {sql_norm} = ? FETCH FIRST 1 ROWS ONLY", tractora_norm),
                (f"SELECT CODTRA, CODCHF FROM {TABLE_NAME} WHERE {sql_norm} LIKE ? FETCH FIRST 1 ROWS ONLY", f"%{tractora_norm}%"),
                (f"SELECT CODTRA, CODCHF FROM {TABLE_NAME} WHERE UPPER(CODCAM) LIKE ? FETCH FIRST 1 ROWS ONLY", f"%{tractora_norm}%"),
            ]
            row = None
            for q, param in queries:
                cur.execute(q, param)
                row = cur.fetchone()
                if row:
                    break
            if row:
                cif = _safe_str(getattr(row, "CODTRA", ""))      # CIF
                agencia = _safe_str(getattr(row, "CODCHF", ""))  # Agencia
        finally:
            try: conn.close()
            except Exception: pass
    except Exception:
        # Falla silenciosa, intentamos fallback
        pass

    # Intento 2: Excel CHF
    if not cif or not agencia:
        try:
            df_chf = _load_chf_table()
            if not df_chf.empty:
                m = df_chf[df_chf["TRACTORA"] == tractora_norm]
                if not m.empty:
                    if not cif:
                        cif = _safe_str(m.iloc[0]["CIF"])
                    if not agencia:
                        agencia = _safe_str(m.iloc[0]["AGENCIA"])
        except Exception:
            pass

    _odbc_cache[tractora_norm] = (cif, agencia)
    return cif, agencia


# ─── QR ─────────────────────────────────────────────────────────────────
def make_qr_png(data: str) -> bytes:
    qr = qrcode.QRCode(box_size=10, border=2)
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


# ─── Code39 ─────────────────────────────────────────────────────────────
def make_barcode_png(value: str, module_height: float = 8.0) -> bytes | None:
    """Genera PNG del código Code128."""
    value = _safe_str(value)
    if not value:
        return None
    try:
        import barcode
        from barcode.writer import ImageWriter
    except ImportError:
        return None
    Code128 = barcode.get_barcode_class("code128")
    code = Code128(value, writer=ImageWriter())
    bio = io.BytesIO()
    code.write(bio, options={
        "module_width": 0.28,
        "module_height": module_height,
        "font_size": 0,
        "text_distance": 0,
        "quiet_zone": 2.0,
        "write_text": False,
        "background": "white",
        "foreground": "black",
    })
    return bio.getvalue()

# Alias para compatibilidad
make_code39_png = make_barcode_png


# ─── Word export (diseño PRO003 — una sola hoja A4) ─────────────────────
def export_word(payload: dict, destino: str = "", precintos: list[dict] | None = None, meta: dict | None = None) -> Path:
    """
    Genera el Word con el diseño PRO003 (idéntico al script Python actual):
      - Cabecera con subtítulo + ciudad GRANDE + agencia/CIF/generado a la derecha
      - Bloque QR + 6 cards (TRACTORA/REMOLQUE/Nº/FECHA/CIF/AGENCIA)
      - Grid de tarjetas de precintos 2 columnas con código de barras Code39
    Devuelve la ruta del .docx generado en SAVE_DIR/.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    precintos = precintos or []
    n_prec = len(precintos)

    SAVE_DIR.mkdir(parents=True, exist_ok=True)
    D_raw = str(payload.get("D", "FECHA"))
    N_val = str(payload.get("N", "NNN"))
    dest_slug = _slug(destino)
    base = f"QR_{D_raw}_{N_val}" + (f"_{dest_slug}" if dest_slug else "")
    path = SAVE_DIR / f"{base}.docx"

    doc = Document()

    # A4 vertical + márgenes estrechos
    for section in doc.sections:
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.top_margin = Inches(0.25)
        section.bottom_margin = Inches(0.20)
        section.left_margin = Inches(0.40)
        section.right_margin = Inches(0.40)

    # Normal sin espaciado extra
    try:
        normal = doc.styles["Normal"]
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.line_spacing = 1.0
    except Exception:
        pass

    # Escalas fijas (calibradas para que 12 precintos entren en 1 hoja A4)
    SC = dict(city_pt=36, sub_pt=9, agency_pt=12, cif_pt=9, gen_pt=9,
              qr_w=1.2, data_lbl=8, data_val=11,
              seal_h=12, seal_c=9, centro_pt=8, pos_pt=8,
              num_pt=10, bc_w=2.3, bc_mh=4.5,
              sp_before_seals=2, card_before=1, card_after=1, div_after=3)

    # Helpers
    def _remove_borders(tbl):
        tblPr = tbl._element.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl._element.insert(0, tblPr)
        tblB = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "nil")
            tblB.append(e)
        tblPr.append(tblB)

    def _set_borders(cell, sides=("top", "bottom", "left", "right"),
                     val="single", sz="4", color="AAAAAA"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in sides:
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), val)
            b.set(qn("w:sz"), sz)
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), color)
            tcBorders.append(b)
        tcPr.append(tcBorders)

    def _set_cell_margins(cell, twips=40):
        tc_props = cell._tc.get_or_add_tcPr()
        tc_margins = OxmlElement("w:tcMar")
        for side in ("top", "left", "bottom", "right"):
            m = OxmlElement(f"w:{side}")
            m.set(qn("w:w"), str(twips))
            m.set(qn("w:type"), "dxa")
            tc_margins.append(m)
        tc_props.append(tc_margins)

    def _clear(cell):
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)

    def _force_font(run, name):
        run.font.name = name
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts.set(qn(k), name)

    def _set_spacing(run, pt_value):
        rPr = run._element.get_or_add_rPr()
        existing = rPr.find(qn("w:spacing"))
        if existing is not None:
            rPr.remove(existing)
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(int(pt_value * 20)))
        rPr.append(sp)

    def _add_para(container, text="", bold=False, size=10, color=None,
                  align=None, after=2, before=0, italic=False, font=None):
        p = container.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(before)
        if align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
        if font:
            _force_font(r, font)
        return p, r

    def _tab_right(para, pos_inches=7.4):
        pPr = para._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(int(pos_inches * 1440)))
        tabs.append(tab)
        pPr.append(tabs)

    # Datos
    agencia = str(payload.get("E", ""))
    cif_val = str(payload.get("C", ""))
    D_display = f"{D_raw[:4]}-{D_raw[4:6]}-{D_raw[6:]}" if (len(D_raw) == 8 and D_raw.isdigit()) else D_raw
    gen_dt = datetime.now().strftime("%Y-%m-%d %H:%M")

    # ─── 1. Cabecera ────────────────────────────────────────────
    hdr = doc.add_table(rows=1, cols=2)
    hdr.autofit = False
    hdr.columns[0].width = Inches(4.7)
    hdr.columns[1].width = Inches(2.8)
    _remove_borders(hdr)

    cl = hdr.cell(0, 0); _clear(cl)
    _, r_sub = _add_para(cl, "CARGA TEKU · CÓDIGO BLEECKER",
                         size=SC["sub_pt"], color="7A7A7A",
                         after=2, bold=True, font="Helvetica")
    _set_spacing(r_sub, 2.2)

    # Auto-shrink ciudad
    dest_text = (destino or "—").upper()
    L = len(dest_text)
    city_pt = SC["city_pt"] if L <= 8 else (34 if L <= 12 else (28 if L <= 16 else (22 if L <= 22 else 18)))
    p_city, _ = _add_para(cl, dest_text, bold=True, size=city_pt, color="111111",
                          after=2, font="Helvetica")
    p_city.paragraph_format.line_spacing = 0.95

    cr = hdr.cell(0, 1); _clear(cr)
    cr.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for txt, b, sz, col in [(agencia, True, SC["agency_pt"], "111111"),
                            (f"CIF {cif_val}", False, SC["cif_pt"], "666666"),
                            (f"Generado {gen_dt}", False, SC["gen_pt"], "888888")]:
        _add_para(cr, txt, bold=b, size=sz, color=col,
                  align="right", after=2, font="Helvetica")

    # Divisoria
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(SC["div_after"])
    p_div.paragraph_format.space_before = Pt(2)
    pPr_div = p_div._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "CFCFCF")
    pBdr.append(bot); pPr_div.append(pBdr)

    # ─── 2. QR + tabla datos ────────────────────────────────────
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    info = doc.add_table(rows=1, cols=2)
    info.autofit = False
    info.columns[0].width = Inches(1.45)
    info.columns[1].width = Inches(6.05)
    # IMPORTANTE: hay que aplicar el ancho también a las celdas individuales
    # para que Word respete las columnas (bug conocido de python-docx)
    info.cell(0, 0).width = Inches(1.45)
    info.cell(0, 1).width = Inches(6.05)
    # Fijar altura mínima para que ambas celdas (QR y datos) terminen alineadas
    info.rows[0].height = Inches(1.5)
    info.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    _remove_borders(info)

    qr_cell = info.cell(0, 0); _clear(qr_cell)
    qr_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_borders(qr_cell)
    _set_cell_margins(qr_cell, twips=60)

    compact = json.dumps(payload, separators=(",", ":"), ensure_ascii=False)
    qr_png = make_qr_png(compact)
    p_qr = qr_cell.add_paragraph()
    p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_qr.paragraph_format.space_after = Pt(1)
    p_qr.paragraph_format.space_before = Pt(2)
    p_qr.add_run().add_picture(io.BytesIO(qr_png), width=Inches(1.3))

    p_cap = qr_cell.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(1)
    r_cap = p_cap.add_run("QR - TEKU/BLEECKER")
    r_cap.font.size = Pt(7)
    r_cap.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    _force_font(r_cap, "Helvetica")
    _set_spacing(r_cap, 1.0)

    data_cell = info.cell(0, 1); _clear(data_cell)
    data_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_borders(data_cell)
    _set_cell_margins(data_cell, twips=60)

    sub = data_cell.add_table(rows=4, cols=2)
    sub.autofit = False
    sub.columns[0].width = Inches(2.85)
    sub.columns[1].width = Inches(2.85)
    _remove_borders(sub)

    meta = meta or {}
    playa_val = _safe_str(meta.get("playa", ""))
    muelle_val = _safe_str(meta.get("muelle", ""))
    data_rows = [
        ("TRACTORA",  str(payload.get("T", "")), "REMOLQUE", str(payload.get("R", ""))),
        ("Nº CAMIÓN", str(payload.get("N", "")), "FECHA",    D_display),
        ("CIF",       cif_val,                    "AGENCIA",  agencia),
        ("MUELLE",    muelle_val,                  "PLAYA",    playa_val),
    ]
    for ri, (l1, v1, l2, v2) in enumerate(data_rows):
        cells_row = sub.rows[ri].cells
        for ci, (lbl, val) in enumerate([(l1, v1), (l2, v2)]):
            c = cells_row[ci]; _clear(c)
            _set_borders(c)
            p_lbl = c.add_paragraph()
            p_lbl.paragraph_format.space_after = Pt(1)
            p_lbl.paragraph_format.space_before = Pt(4)
            r_lbl = p_lbl.add_run(lbl)
            r_lbl.font.size = Pt(SC["data_lbl"])
            r_lbl.font.color.rgb = RGBColor(0x9A, 0x9A, 0x9A)
            _force_font(r_lbl, "Helvetica")
            _set_spacing(r_lbl, 1.5)
            p_val = c.add_paragraph()
            p_val.paragraph_format.space_after = Pt(4)
            r_val = p_val.add_run(val)
            r_val.bold = True
            r_val.font.size = Pt(SC["data_val"])
            r_val.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
            _force_font(r_val, "Helvetica")

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(SC["sp_before_seals"])

    # ─── 3. Precintos ───────────────────────────────────────────
    if n_prec > 0:
        p_ph = doc.add_paragraph()
        p_ph.paragraph_format.space_after = Pt(6)
        r_ph1 = p_ph.add_run("Precintos")
        r_ph1.bold = True
        r_ph1.font.size = Pt(SC["seal_h"])
        _force_font(r_ph1, "Helvetica")
        p_ph.add_run("\t")
        r_ph2 = p_ph.add_run(f"{n_prec} PRECINTO{'S' if n_prec != 1 else ''}")
        r_ph2.font.size = Pt(SC["seal_c"])
        r_ph2.bold = True
        r_ph2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        _force_font(r_ph2, "Helvetica")
        _tab_right(p_ph, pos_inches=7.4)

        n_rows_grid = (n_prec + 1) // 2
        grid = doc.add_table(rows=n_rows_grid, cols=2)
        grid.autofit = False
        grid.columns[0].width = Inches(3.7)
        grid.columns[1].width = Inches(3.7)
        _remove_borders(grid)

        for idx, pdi in enumerate(precintos):
            centro = _safe_str(pdi.get("centro", ""))
            prec_val = _safe_str(pdi.get("precinto", ""))
            ri, ci = idx // 2, idx % 2
            card = grid.cell(ri, ci); _clear(card)
            _set_borders(card)
            card.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            p_hdr = card.add_paragraph()
            p_hdr.paragraph_format.space_after = Pt(1)
            p_hdr.paragraph_format.space_before = Pt(SC["card_before"])
            r_c = p_hdr.add_run(centro.upper() if centro else "CENTRO")
            r_c.font.size = Pt(SC["centro_pt"])
            r_c.bold = True
            r_c.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
            _force_font(r_c, "Helvetica")
            _set_spacing(r_c, 1.8)
            p_hdr.add_run("\t")
            r_cnt = p_hdr.add_run(f"{idx+1:02d} / {n_prec:02d}")
            r_cnt.font.size = Pt(SC["pos_pt"])
            r_cnt.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            _force_font(r_cnt, "Helvetica")
            _tab_right(p_hdr, pos_inches=3.55)

            p_num = card.add_paragraph()
            p_num.paragraph_format.space_after = Pt(1)
            r_num = p_num.add_run(prec_val)
            r_num.bold = True
            r_num.font.size = Pt(SC["num_pt"])
            _force_font(r_num, "Helvetica")

            bc_png = make_barcode_png(prec_val, module_height=SC["bc_mh"])
            if bc_png:
                p_bc = card.add_paragraph()
                p_bc.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_bc.paragraph_format.space_after = Pt(SC["card_after"])
                p_bc.add_run().add_picture(io.BytesIO(bc_png), width=Inches(SC["bc_w"]))

    # ─── 4. Pie ─────────────────────────────────────────────────
    p_foot = doc.add_paragraph()
    p_foot.paragraph_format.space_before = Pt(SC["sp_before_seals"])
    p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr_f = p_foot._p.get_or_add_pPr()
    pBdr_f = OxmlElement("w:pBdr")
    top_b = OxmlElement("w:top")
    top_b.set(qn("w:val"), "single"); top_b.set(qn("w:sz"), "2")
    top_b.set(qn("w:space"), "6"); top_b.set(qn("w:color"), "E3E3E3")
    pBdr_f.append(top_b); pPr_f.append(pBdr_f)
    r_f = p_foot.add_run(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    r_f.font.size = Pt(7)
    r_f.italic = True
    r_f.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    _force_font(r_f, "Helvetica")

    doc.save(str(path))
    return path


# ─── Impresión ──────────────────────────────────────────────────────────
def print_file(path: Path | str) -> None:
    p = str(path)
    if platform.system().lower().startswith("win"):
        os.startfile(p, "print")
    elif shutil.which("lpr"):
        subprocess.Popen(["lpr", p])
    else:
        raise RuntimeError("Impresión automática no disponible en este sistema.")
